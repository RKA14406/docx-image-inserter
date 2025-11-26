import os
import uuid
import json
import html
from typing import List

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Inches
import httpx


# ================== إعداد التطبيق ==================

app = FastAPI()

# في الإنتاج يُفضّل تقييد origins لدومين موقعك، لكن هنا نجعلها مفتوحة للتجربة
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SESSIONS_DIR = os.path.join(BASE_DIR, "sessions")
os.makedirs(SESSIONS_DIR, exist_ok=True)

# مفاتيح Groq من الـ environment في السيرفر (لا تكتبها في الكود!)
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.1-8b-instant"  # غيّر الموديل إذا أحببت


# ================== دوال مساعدة ==================

def create_session_dir() -> str:
    session_id = str(uuid.uuid4())
    session_path = os.path.join(SESSIONS_DIR, session_id)
    os.makedirs(session_path, exist_ok=True)
    return session_id


def docx_to_simple_html(doc_path: str):
    """
    تحويل بسيط لملف وورد إلى HTML:
    - كل فقرة تصبح <p data-para-id="p_i">...</p>
    - النص يُعمل له escape عشان علامات HTML.
    هذا ليس نفس تنسيق الوورد الحقيقي، لكنه كافي لاختيار الأسطر.
    """
    doc = Document(doc_path)
    paragraphs = []
    html_parts = []

    for i, p in enumerate(doc.paragraphs):
        raw_text = p.text or ""
        display_text = raw_text.strip()
        if not display_text:
            display_text = "‎"  # حرف غير مرئي حتى لا تكون الفقرة فارغة

        para_id = f"p_{i}"
        paragraphs.append({"id": para_id, "text": raw_text})
        safe_text = html.escape(display_text)
        html_parts.append(f'<p data-para-id="{para_id}">{safe_text}</p>')

    html_str = "\n".join(html_parts)
    return html_str, paragraphs


# ================== نماذج للـ API ==================

class Placement(BaseModel):
    image_id: str
    para_id: str
    position: str  # "after" أو "before"


class ApplyRequest(BaseModel):
    session_id: str
    placements: List[Placement]


class AISuggestRequest(BaseModel):
    session_id: str
    instructions: str  # تعليمات الوالد بالكلام العادي (العربية)


# ================== API: رفع الملف والصور ==================

@app.post("/api/session/create")
async def create_session(
    docx_file: UploadFile = File(...),
    images: List[UploadFile] = File(default=[]),
):
    """
    يرفع ملف وورد + صور
    يرجع:
      - session_id
      - HTML بسيط للمعاينة
      - قائمة الفقرات
      - قائمة الصور
    """
    session_id = create_session_dir()
    session_path = os.path.join(SESSIONS_DIR, session_id)

    # حفظ ملف الوورد
    docx_path = os.path.join(session_path, "original.docx")
    with open(docx_path, "wb") as f:
        f.write(await docx_file.read())

    # حفظ الصور
    image_infos = []
    for img in images:
        ext = os.path.splitext(img.filename)[1] or ".png"
        img_id = str(uuid.uuid4())
        stored_name = f"{img_id}{ext}"
        img_path = os.path.join(session_path, stored_name)
        with open(img_path, "wb") as f:
            f.write(await img.read())

        image_infos.append(
            {
                "id": img_id,
                "filename": img.filename,
                "stored_name": stored_name,
            }
        )

    # تحويل الملف إلى HTML بسيط
    html_preview, paragraphs = docx_to_simple_html(docx_path)

    # تخزين البيانات الوصفية
    meta = {
        "docx_path": "original.docx",
        "images": image_infos,
        "paragraphs": paragraphs,
    }
    with open(os.path.join(session_path, "meta.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    return JSONResponse(
        {
            "session_id": session_id,
            "html_preview": html_preview,
            "images": image_infos,
            "paragraphs": paragraphs,
        }
    )


# ================== API: تطبيق الأماكن وإرجاع ملف وورد ==================

@app.post("/api/session/apply")
def apply_placements(body: ApplyRequest):
    """
    يأخذ session_id + قائمة أماكن الصور
    يرجع ملف وورد النهائي للتحميل
    """
    session_id = body.session_id
    session_path = os.path.join(SESSIONS_DIR, session_id)
    if not os.path.isdir(session_path):
        return JSONResponse({"error": "Invalid session_id"}, status_code=400)

    meta_path = os.path.join(session_path, "meta.json")
    if not os.path.exists(meta_path):
        return JSONResponse({"error": "Session metadata not found"}, status_code=400)

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    docx_path = os.path.join(session_path, meta["docx_path"])
    doc = Document(docx_path)

    # خريطة para_id -> index الأصلي
    para_index_by_id = {p["id"]: idx for idx, p in enumerate(meta["paragraphs"])}

    # خريطة image_id -> بيانات الصورة
    image_by_id = {img["id"]: img for img in meta["images"]}

    # نرتّب من الأسفل للأعلى عشان ما نخرب الـ indices
    placements_sorted = sorted(
        body.placements,
        key=lambda p: para_index_by_id.get(p.para_id, -1),
        reverse=True,
    )

    for placement in placements_sorted:
        if placement.image_id not in image_by_id:
            continue
        if placement.para_id not in para_index_by_id:
            continue

        img_info = image_by_id[placement.image_id]
        img_path = os.path.join(session_path, img_info["stored_name"])
        idx = para_index_by_id[placement.para_id]

        if idx < 0 or idx >= len(doc.paragraphs):
            continue

        target_para = doc.paragraphs[idx]

        # نضيف الصورة كفقرة جديدة في نهاية المستند ثم ننقلها قبل / بعد الفقرة الهدف
        doc.add_picture(img_path, width=Inches(3))
        pic_para = doc.paragraphs[-1]

        target_el = target_para._p
        pic_el = pic_para._p

        if placement.position == "before":
            target_el.addprevious(pic_el)
        else:  # "after"
            target_el.addnext(pic_el)

    output_path = os.path.join(session_path, "output.docx")
    doc.save(output_path)

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="modified.docx",
    )


# ================== API: اقتراح أماكن باستخدام Groq ==================

@app.post("/api/session/suggest-placements")
async def suggest_placements(body: AISuggestRequest):
    """
    يأخذ:
      - session_id
      - instructions (تعليمات عربية)

    يرجع:
      - ai_placements: قائمة اقتراحات مثل:
        [
          {"image_label": "الشعار", "para_id": "p_0", "position": "before"},
          ...
        ]
    """
    if not GROQ_API_KEY:
        return JSONResponse(
            {"error": "GROQ_API_KEY not configured on server"},
            status_code=500,
        )

    session_path = os.path.join(SESSIONS_DIR, body.session_id)
    if not os.path.isdir(session_path):
        return JSONResponse({"error": "Invalid session_id"}, status_code=400)

    meta_path = os.path.join(session_path, "meta.json")
    if not os.path.exists(meta_path):
        return JSONResponse({"error": "Session metadata not found"}, status_code=400)

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    paragraphs = meta.get("paragraphs", [])
    outline_paras = [
        {
            "id": p["id"],
            "sample": (p.get("text") or "")[:120],
        }
        for p in paragraphs
    ]

    system_prompt = """
You are a helper that converts natural Arabic instructions into JSON rules
for inserting images into a Word document.

You receive:
- A list of paragraphs with ids and text samples.
- Natural language instructions in Arabic from the user describing where to insert each image.
- The user does NOT know programming or technical terms.

You MUST output ONLY valid JSON, no explanations, with this exact structure:

[
  {
    "image_label": "اسم تقريبي للصورة (مثلاً 'الشعار' أو 'الرسم البياني')",
    "para_id": "p_5",
    "position": "after"
  }
]

Rules:
- Choose para_id based ONLY on the paragraph samples you have.
- "position" must be either "before" or "after".
- If you are not sure where something goes, make your best guess.
- Do NOT add any other keys.
"""

    user_content = {
        "outline": outline_paras,
        "instructions": body.instructions,
    }

    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": json.dumps(user_content, ensure_ascii=False),
            },
        ],
        "temperature": 0.2,
    }

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json",
    }

    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(GROQ_API_URL, headers=headers, json=payload)
            resp.raise_for_status()
            data = resp.json()
    except httpx.HTTPError as e:
        return JSONResponse(
            {"error": "Error calling Groq API", "details": str(e)}, status_code=500
        )

    content = data["choices"][0]["message"]["content"]

    try:
        ai_placements = json.loads(content)
    except Exception:
        return JSONResponse(
            {"error": "Groq response was not valid JSON", "raw": content},
            status_code=500,
        )

    return JSONResponse({"ai_placements": ai_placements})


# ملاحظة: على Railway / Render ستستخدم أمر تشغيل مثل:
# uvicorn main:app --host 0.0.0.0 --port $PORT
